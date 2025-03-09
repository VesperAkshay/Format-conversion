import os
import asyncio
import subprocess
from typing import List, Optional

from app.utils.base_converter import BaseConverter

class DocumentConverter(BaseConverter):
    """
    Converter for document file formats.
    Supports conversions between various document formats like pdf, docx, txt, etc.
    """
    
    def __init__(self):
        super().__init__()
        
        # Define supported formats
        self._input_formats = ["pdf", "docx", "doc", "txt", "rtf", "odt", "html", "md"]
        self._output_formats = ["pdf", "docx", "txt", "html", "md"]
    
    async def convert(
        self, 
        file_path: str, 
        target_format: str,
        output_filename: Optional[str] = None
    ) -> str:
        """
        Convert a document file to the specified format.
        
        Args:
            file_path: Path to the file to convert
            target_format: Format to convert to
            output_filename: Optional custom filename for the output file
            
        Returns:
            Path to the converted file
        
        Raises:
            ValueError: If the input format or target format is not supported
            Exception: If the conversion fails
        """
        input_format = self._get_file_extension(file_path)
        self._validate_formats(input_format, target_format)
        
        output_path = self._generate_output_path(file_path, target_format, output_filename)
        
        # Perform the conversion based on input and output formats
        if input_format == "pdf" and target_format == "docx":
            await self._pdf_to_docx(file_path, output_path)
        elif input_format == "pdf" and target_format == "txt":
            await self._pdf_to_txt(file_path, output_path)
        elif input_format == "pdf" and target_format == "html":
            await self._pdf_to_html(file_path, output_path)
        elif input_format == "docx" and target_format == "pdf":
            await self._docx_to_pdf(file_path, output_path)
        elif input_format == "docx" and target_format == "txt":
            await self._docx_to_txt(file_path, output_path)
        elif input_format == "docx" and target_format == "html":
            await self._docx_to_html(file_path, output_path)
        elif input_format == "txt" and target_format == "pdf":
            await self._txt_to_pdf(file_path, output_path)
        elif input_format == "txt" and target_format == "docx":
            await self._txt_to_docx(file_path, output_path)
        elif input_format == "txt" and target_format == "html":
            await self._txt_to_html(file_path, output_path)
        elif input_format == "html" and target_format == "pdf":
            await self._html_to_pdf(file_path, output_path)
        elif input_format == "html" and target_format == "docx":
            await self._html_to_docx(file_path, output_path)
        elif input_format == "html" and target_format == "txt":
            await self._html_to_txt(file_path, output_path)
        elif input_format == "md" and target_format == "pdf":
            await self._md_to_pdf(file_path, output_path)
        elif input_format == "md" and target_format == "docx":
            await self._md_to_docx(file_path, output_path)
        elif input_format == "md" and target_format == "html":
            await self._md_to_html(file_path, output_path)
        else:
            raise ValueError(f"Conversion from {input_format} to {target_format} is not supported")
        
        return output_path
    
    def get_supported_input_formats(self) -> List[str]:
        """Get a list of supported input formats"""
        return self._input_formats
    
    def get_supported_output_formats(self) -> List[str]:
        """Get a list of supported output formats"""
        return self._output_formats
    
    # Conversion methods
    
    async def _pdf_to_docx(self, input_path: str, output_path: str) -> None:
        """Convert PDF to DOCX"""
        try:
            from pdf2docx import Converter
            
            cv = Converter(input_path)
            cv.convert(output_path)
            cv.close()
        except ImportError:
            raise Exception("pdf2docx library is required for PDF to DOCX conversion")
    
    async def _pdf_to_txt(self, input_path: str, output_path: str) -> None:
        """Convert PDF to TXT"""
        try:
            import PyPDF2
            
            with open(input_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num in range(len(reader.pages)):
                    text += reader.pages[page_num].extract_text()
            
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(text)
        except ImportError:
            raise Exception("PyPDF2 library is required for PDF to TXT conversion")
    
    async def _pdf_to_html(self, input_path: str, output_path: str) -> None:
        """Convert PDF to HTML"""
        try:
            # This is a placeholder. For a production app, you might want to use a more robust solution
            # like pdf2htmlEX or a similar tool
            import PyPDF2
            
            with open(input_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num in range(len(reader.pages)):
                    text += reader.pages[page_num].extract_text()
            
            # Convert text to HTML
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>Converted PDF</title>
</head>
<body>
    <pre>{text}</pre>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(html_content)
        except ImportError:
            raise Exception("PyPDF2 library is required for PDF to HTML conversion")
    
    async def _docx_to_pdf(self, input_path: str, output_path: str) -> None:
        """Convert DOCX to PDF"""
        try:
            from docx2pdf import convert
            convert(input_path, output_path)
        except ImportError:
            raise Exception("docx2pdf library is required for DOCX to PDF conversion")
    
    async def _docx_to_txt(self, input_path: str, output_path: str) -> None:
        """Convert DOCX to TXT"""
        try:
            import docx
            
            doc = docx.Document(input_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(text)
        except ImportError:
            raise Exception("python-docx library is required for DOCX to TXT conversion")
    
    async def _docx_to_html(self, input_path: str, output_path: str) -> None:
        """Convert DOCX to HTML"""
        try:
            import docx
            
            doc = docx.Document(input_path)
            html_parts = ["<!DOCTYPE html>\n<html>\n<head>\n<title>Converted Document</title>\n</head>\n<body>"]
            
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    html_parts.append(f"<p>{paragraph.text}</p>")
            
            html_parts.append("</body>\n</html>")
            html_content = "\n".join(html_parts)
            
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(html_content)
        except ImportError:
            raise Exception("python-docx library is required for DOCX to HTML conversion")
    
    async def _txt_to_pdf(self, input_path: str, output_path: str) -> None:
        """Convert TXT to PDF"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            import html
            
            # Read the text file
            with open(input_path, 'r', encoding='utf-8', errors='replace') as file:
                text = file.read()
            
            # Create a PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Define styles
            styles = getSampleStyleSheet()
            normal_style = styles['Normal']
            
            # Process the text content
            content = []
            
            # Split text into paragraphs
            paragraphs = text.split('\n\n')
            
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Replace single newlines with spaces within paragraphs
                    paragraph = paragraph.replace('\n', ' ')
                    # Escape any HTML-like characters to prevent rendering issues
                    paragraph = html.escape(paragraph)
                    p = Paragraph(paragraph, normal_style)
                    content.append(p)
                    content.append(Spacer(1, 0.2 * inch))
            
            # Build the PDF
            doc.build(content)
            
        except Exception as e:
            raise Exception(f"Error in TXT to PDF conversion: {str(e)}")
    
    async def _txt_to_docx(self, input_path: str, output_path: str) -> None:
        """Convert TXT to DOCX"""
        try:
            import docx
            
            with open(input_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            doc = docx.Document()
            
            # Split text into paragraphs
            paragraphs = text.split('\n\n')
            
            # Add paragraphs to document
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            doc.save(output_path)
        except ImportError:
            raise Exception("python-docx library is required for TXT to DOCX conversion")
    
    async def _txt_to_html(self, input_path: str, output_path: str) -> None:
        """Convert TXT to HTML"""
        try:
            from bs4 import BeautifulSoup
            
            with open(input_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            soup = BeautifulSoup(text, 'html.parser')
            html_content = soup.get_text()
            
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(html_content)
        except ImportError:
            raise Exception("BeautifulSoup library is required for TXT to HTML conversion")
    
    async def _html_to_pdf(self, input_path: str, output_path: str) -> None:
        """Convert HTML to PDF"""
        try:
            import pdfkit
            
            pdfkit.from_file(input_path, output_path)
        except ImportError:
            # Fallback to a simpler approach if pdfkit is not available
            try:
                from bs4 import BeautifulSoup
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                
                with open(input_path, 'r', encoding='utf-8') as file:
                    html_content = file.read()
                
                soup = BeautifulSoup(html_content, 'html.parser')
                text = soup.get_text()
                
                c = canvas.Canvas(output_path, pagesize=letter)
                width, height = letter
                
                # Split text into lines
                lines = text.split('\n')
                
                # Set font and size
                c.setFont("Helvetica", 12)
                
                # Calculate line height
                line_height = 14
                
                # Start position
                y = height - 50
                
                # Add lines to PDF
                for line in lines:
                    if y < 50:  # Check if we need a new page
                        c.showPage()
                        y = height - 50
                        c.setFont("Helvetica", 12)
                    
                    c.drawString(50, y, line[:80])  # Limit line length
                    y -= line_height
                
                c.save()
            except ImportError:
                raise Exception("BeautifulSoup and reportlab libraries are required for HTML to PDF conversion")
    
    async def _html_to_docx(self, input_path: str, output_path: str) -> None:
        """Convert HTML to DOCX"""
        try:
            from bs4 import BeautifulSoup
            import docx
            
            with open(input_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            text = soup.get_text()
            
            doc = docx.Document()
            
            # Split text into paragraphs
            paragraphs = text.split('\n\n')
            
            # Add paragraphs to document
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            doc.save(output_path)
        except ImportError:
            raise Exception("BeautifulSoup and python-docx libraries are required for HTML to DOCX conversion")
    
    async def _html_to_txt(self, input_path: str, output_path: str) -> None:
        """Convert HTML to TXT"""
        try:
            from bs4 import BeautifulSoup
            
            with open(input_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            text = soup.get_text()
            
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(text)
        except ImportError:
            raise Exception("BeautifulSoup library is required for HTML to TXT conversion")
    
    async def _md_to_pdf(self, input_path: str, output_path: str) -> None:
        """Convert Markdown to PDF"""
        # First convert MD to HTML, then HTML to PDF
        html_path = output_path.replace('.pdf', '.html')
        
        await self._md_to_html(input_path, html_path)
        await self._html_to_pdf(html_path, output_path)
        
        # Clean up temporary HTML file
        if os.path.exists(html_path):
            os.remove(html_path)
    
    async def _md_to_docx(self, input_path: str, output_path: str) -> None:
        """Convert Markdown to DOCX"""
        # First convert MD to HTML, then HTML to DOCX
        html_path = output_path.replace('.docx', '.html')
        
        await self._md_to_html(input_path, html_path)
        await self._html_to_docx(html_path, output_path)
        
        # Clean up temporary HTML file
        if os.path.exists(html_path):
            os.remove(html_path)
    
    async def _md_to_html(self, input_path: str, output_path: str) -> None:
        """Convert Markdown to HTML"""
        try:
            import markdown
            
            with open(input_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            html_content = markdown.markdown(md_content)
            
            # Wrap in HTML document structure
            html_document = f"""<!DOCTYPE html>
<html>
<head>
    <title>Converted Markdown</title>
</head>
<body>
{html_content}
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(html_document)
        except ImportError:
            raise Exception("markdown library is required for Markdown to HTML conversion") 